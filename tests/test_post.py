import json
import os
import pytest
import platform

from dotenv import load_dotenv
from confluent_kafka import KafkaError
from httpx import AsyncClient, ASGITransport
from redis import Redis
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document

from srt.data_base.models import User, Requirements, Resume
from srt.main import app
from srt.config import MAX_CHAR_REQUIREMENTS, MAX_CHAR_RESUME, KEY_NEW_REQUEST, KEY_NEW_RESUME, KEY_NEW_REQUIREMENTS
from tests.conftest import consumer
from tests.conftest import PATH_TO_TEST_DIRECTORY

load_dotenv()  # Загружает переменные из .env
KAFKA_BOOTSTRAP_SERVERS= os.getenv('KAFKA_BOOTSTRAP_SERVERS')
KAFKA_TOPIC_CONSUMER= os.getenv('KAFKA_TOPIC_CONSUMER')
KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA= os.getenv('KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA')
KAFKA_TOPIC_PRODUCER_FOR_AI_HANDLER= os.getenv('KAFKA_TOPIC_PRODUCER_FOR_AI_HANDLER')


def get_pdf_font():
    if platform.system() == 'Windows':
        try:
            # Путь к Arial на Windows
            font_path = os.path.join(os.environ['WINDIR'], 'Fonts', 'arial.ttf')
            pdfmetrics.registerFont(TTFont('ArialUnicode', font_path))
            return "ArialUnicode"
        except:
            return "Helvetica"
    else:
        try:
            # Для Linux/CI используем DejaVu
            pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            return "DejaVuSans"
        except:
            return "Helvetica"

async def comparison_resume_data(resume: str, data_response: dict, db_session:  AsyncSession, redis_session: Redis):
    # Данные с БД
    result_db = await db_session.execute(select(Resume).where(Resume.resume_id == data_response['resume_id']))
    data_db = result_db.scalar_one_or_none()

    # Данные с Kafka
    data_kafka = None
    for i in range(40):
        try:
            msg = consumer.poll(timeout=1.0)
            if msg is None:
                if i == 39:
                    raise Exception("Не удалось получить сообщение от Kafka!")
                continue
            if msg.key().decode('utf-8') == KEY_NEW_RESUME:
                data_kafka = json.loads(msg.value().decode('utf-8'))
            else:
                raise Exception(f"Ожидался ключ {KEY_NEW_RESUME}, получен {msg.key().decode('utf-8')}")
            break
        except KafkaError as e:
            raise Exception(f"Ошибка Kafka: {e}")

    # Данные из Redis
    data_redis = await redis_session.get(f'resume:{data_response["resume_id"]}')

    # Проверки
    assert data_response["resume_id"] == data_db.resume_id == data_kafka['resume_id']
    assert data_db.user_id == data_kafka['user_id']
    assert data_db.resume == data_kafka['resume'] == data_redis == resume # Проверка данных у резюме

async def comparison_requirements_data(requirements: str, data_response: dict, db_session:  AsyncSession, redis_session: Redis):
    # данные с БД
    result_db = await db_session.execute(select(Requirements).where(Requirements.requirements_id == data_response['requirements_id']))
    data_db = result_db.scalar_one_or_none()

    # данные с kafka
    data_kafka = None
    for i in range(
            40):  # такой большой тайминг, ибо при пересоздании топика может быть большая задержка у первого сообщения
        try:
            msg = consumer.poll(timeout=1.0)
            if msg is None:  # если не успели отослать сообщение, а уже пытаемся его прочитать
                if i == 39:
                    raise Exception("НЕ смогли получить сообщение от kafka!")
                else:
                    continue

            if msg.key().decode('utf-8') == KEY_NEW_REQUIREMENTS:
                data_kafka = json.loads(msg.value().decode('utf-8'))
            else:
                raise Exception(f"Ожидался ключ '{KEY_NEW_REQUIREMENTS}', получен {msg.key().decode('utf-8')}")
            break
        except KafkaError as e:
            raise Exception(f"Ошибка Kafka: {e}")

    data_redis = await redis_session.get(f'requirements:{data_response['requirements_id']}')

    assert data_response["requirements_id"] == data_db.requirements_id == data_kafka['requirements_id']  # проверка на requirements_id
    assert data_db.user_id == data_kafka['user_id']  # проверка на user_id
    assert data_db.requirements == data_kafka['requirements'] == data_redis == requirements  # проверка на сами требования

class TestCreateRequirementsText:
    @pytest.mark.asyncio
    async def test_code_200(self, db_session, clearing_kafka, redis_session, create_user):
        consumer.subscribe([KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA]) # подписка на топик
        requirements_content = 'это требования к резюме'

        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:

            response = await ac.post(
                "/create_requirements/text",
                json={'requirements': requirements_content},
                headers={"Authorization": f"Bearer {create_user['access_token']}"}
            )

            assert response.status_code == 200
            data_response = response.json()# данные ответа сервера

            await comparison_requirements_data(requirements_content, data_response, db_session, redis_session)

    async def test_code_413(self, db_session, clearing_kafka, redis_session, create_user):
        # подписка на топик
        consumer.subscribe([KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA])
        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            data_request = {'requirements': 'требования к резюме которое должно превысить допустимую длину'}
            while True:
                data_request['requirements'] += data_request['requirements']
                if len(data_request['requirements']) > MAX_CHAR_REQUIREMENTS:
                    response = await ac.post(
                        "/create_requirements/text",
                        json=data_request,
                        headers={"Authorization": f"Bearer {create_user['access_token']}"}
                    )
                    assert response.status_code == 413
                    break

class TestCreateRequirementsFile:
    @pytest.mark.asyncio
    @pytest.mark.parametrize(
        'extension, status_code',
        [
            ('.txt', 200),
            ('.pdf', 200),
            ('.docx', 200),
            ('.doc', 400),  # неподдерживаемый формат
        ]
    )
    async def test_send_file(self, extension,status_code, db_session, clearing_kafka, redis_session, create_user):
        consumer.subscribe([KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA])

        # Создаем временный файл для теста
        requirements_content = 'это требования к резюме'
        test_file_path = f"{PATH_TO_TEST_DIRECTORY}/test_requirements" + extension
        if extension == '.txt':
            with open(test_file_path, "w", encoding="utf-8") as f:
                f.write(requirements_content)
        elif extension == '.pdf':
            font_name = get_pdf_font()
            c = canvas.Canvas(test_file_path, pagesize=letter)
            c.setFont(font_name, 12)  # Используем Unicode-шрифт
            c.drawString(100, 750, requirements_content)
            c.save()
        elif extension == '.docx':
            doc = Document() # Создаем новый документ
            doc.add_paragraph(requirements_content)# Добавляем абзац
            doc.save(test_file_path)
        elif extension == '.doc':
            doc = Document()# Создаем новый документ
            doc.add_heading('Мой документ Word', level=1)
            doc.add_paragraph(requirements_content)
            doc.save(test_file_path)

        try:
            async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
            ) as ac:
                # Открываем файл и отправляем его как часть multipart-формы
                with open(test_file_path, "rb") as f:
                    response = await ac.post(
                        "/create_requirements/file",
                        files={"file": (test_file_path, f, "text/plain")},
                        headers={"Authorization": f"Bearer {create_user['access_token']}"}
                    )

                if status_code == 200:
                    assert response.status_code == status_code
                    data_response = response.json()  # Данные ответа сервера
                    await comparison_requirements_data(requirements_content, data_response, db_session, redis_session)
                elif status_code == 400:
                    assert response.status_code == status_code
        finally:
            # Удаляем временный файл
            if os.path.exists(test_file_path):
                os.remove(test_file_path)

    @pytest.mark.asyncio
    async def test_code_400(self, create_user):
        test_file_path = f"{PATH_TO_TEST_DIRECTORY}/test_requirements.txt"
        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            try:
                with open(test_file_path, "w", encoding="utf-8") as f:
                    pass # в файл ничего не записываем

                # Открываем файл и отправляем его как часть multipart-формы
                with open(test_file_path, "rb") as f:
                    response = await ac.post(
                        "/create_requirements/file",
                        files={"file": (test_file_path, f, "text/plain")},
                        headers={"Authorization": f"Bearer {create_user['access_token']}"}
                    )
                assert response.status_code == 400
            finally:
                # Удаляем временный файл
                if os.path.exists(test_file_path):
                    os.remove(test_file_path)

    @pytest.mark.asyncio
    async def test_code_413(self, create_user):
        resume_content = 'резюме которое должно превысить допустимую длину'
        test_file_path = f"{PATH_TO_TEST_DIRECTORY}/test_requirements.txt"

        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            try:

                while True:
                    resume_content += resume_content

                    if len(resume_content) > MAX_CHAR_RESUME:
                        with open(test_file_path, "w", encoding="utf-8") as f:
                            f.write(resume_content)

                        # Открываем файл и отправляем его как часть multipart-формы
                        with open(test_file_path, "rb") as f:
                            response = await ac.post(
                                "/create_requirements/file",
                                files={"file": (test_file_path, f, "text/plain")},
                                headers={"Authorization": f"Bearer {create_user['access_token']}"}
                            )
                        assert response.status_code == 413
                        break
            finally:
                # Удаляем временный файл
                if os.path.exists(test_file_path):
                    os.remove(test_file_path)

class TestCreateResumeText:
    @pytest.mark.asyncio
    async def test_code_200(self, db_session, clearing_kafka, redis_session, create_user):
        # подписка на топик
        consumer.subscribe([KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA])
        resume_content = 'это резюме'

        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:

            response = await ac.post(
                "/create_resume/text",
                json={'resume': resume_content},
                headers={"Authorization": f"Bearer {create_user['access_token']}"}
            )

            assert response.status_code == 200
            data_response = response.json()# данные ответа сервера

            await comparison_resume_data(resume_content, data_response, db_session, redis_session)

    async def test_code_413(self, create_user):
        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            data_request = {'resume': 'резюме которое должно превысить допустимую длину'}
            while True:
                data_request['resume'] += data_request['resume']

                if len(data_request['resume']) > MAX_CHAR_RESUME:
                    response = await ac.post(
                        "/create_resume/text",
                        json=data_request,
                        headers={"Authorization": f"Bearer {create_user['access_token']}"}
                    )
                    assert response.status_code == 413
                    break

class TestCreateResumeFile:
    @pytest.mark.asyncio
    @pytest.mark.parametrize(
        'extension, status_code',
        [
            ('.txt', 200),
            ('.pdf', 200),
            ('.docx', 200),
            ('.doc', 400),  # неподдерживаемый формат
        ]
    )
    async def test_send_file(self, extension, status_code, db_session, clearing_kafka, redis_session, create_user):
        consumer.subscribe([KAFKA_TOPIC_PRODUCER_FOR_UPLOADING_DATA])

        # Создаем временный файл для теста
        resume_content = "Это тестовое резюме в файле."
        test_file_path = f"{PATH_TO_TEST_DIRECTORY}/test_resume" + extension
        if extension == '.txt':
            with open(test_file_path, "w", encoding="utf-8") as f:
                f.write(resume_content)
        elif extension == '.pdf':
            font_name = get_pdf_font()
            c = canvas.Canvas(test_file_path, pagesize=letter)
            c.setFont(font_name, 12)  # Используем Unicode-шрифт
            c.drawString(100, 750, resume_content)
            c.save()
        elif extension == '.docx':
            doc = Document() # Создаем новый документ
            doc.add_paragraph(resume_content)# Добавляем абзац
            doc.save(test_file_path)
        elif extension == '.doc':
            doc = Document()# Создаем новый документ
            doc.add_heading('Мой документ Word', level=1)
            doc.add_paragraph(resume_content)
            doc.save(test_file_path)

        try:
            async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
            ) as ac:
                # Открываем файл и отправляем его как часть multipart-формы
                with open(test_file_path, "rb") as f:
                    response = await ac.post(
                        "/create_resume/file",
                        files={"file": (test_file_path, f, "text/plain")},
                        headers={"Authorization": f"Bearer {create_user['access_token']}"}
                    )
                if status_code == 200:
                    assert response.status_code == status_code
                    data_response = response.json()# Данные ответа сервера
                    await comparison_resume_data(resume_content, data_response, db_session, redis_session)
                elif status_code == 400:
                    assert response.status_code == status_code
        finally:
            # Удаляем временный файл
            if os.path.exists(test_file_path):
                os.remove(test_file_path)


    @pytest.mark.asyncio
    async def test_code_400(self, create_user):
        test_file_path = f"{PATH_TO_TEST_DIRECTORY}/test_resume.txt"
        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            try:
                with open(test_file_path, "w", encoding="utf-8") as f:
                    pass # в файл ничего не записываем

                # Открываем файл и отправляем его как часть multipart-формы
                with open(test_file_path, "rb") as f:
                    response = await ac.post(
                        "/create_resume/file",
                        files={"file": (test_file_path, f, "text/plain")},
                        headers={"Authorization": f"Bearer {create_user['access_token']}"}
                    )
                assert response.status_code == 400
            finally:
                # Удаляем временный файл
                if os.path.exists(test_file_path):
                    os.remove(test_file_path)

    @pytest.mark.asyncio
    async def test_code_413(self, create_user):
        resume_content = 'резюме которое должно превысить допустимую длину'
        test_file_path = f"{PATH_TO_TEST_DIRECTORY}/test_resume.txt"

        try:

            while len(resume_content) <= MAX_CHAR_RESUME:
                resume_content += resume_content

            with open(test_file_path, "w", encoding="utf-8") as f:
                f.write(resume_content)

                async with AsyncClient(
                        transport=ASGITransport(app),
                        base_url="http://test",
                ) as ac:
                    with open(test_file_path, "rb") as f:
                        response = await ac.post(
                            "/create_resume/file",
                            files={"file": (test_file_path, f, "text/plain")},
                            headers={"Authorization": f"Bearer {create_user['access_token']}"}
                        )
                    assert response.status_code == 413
        finally:
            # Удаляем временный файл
            if os.path.exists(test_file_path):
                os.remove(test_file_path)

class TestStartProcessing:
    @pytest.mark.asyncio
    async def test_code_200(self, db_session, clearing_kafka, redis_session, create_requirements_and_resume):
        consumer.subscribe([KAFKA_TOPIC_PRODUCER_FOR_AI_HANDLER]) # подписка на топик
        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            callback_url = 'http://test_url/'
            response = await ac.post(
                "/start_processing",
                json={
                    "requirements_id": create_requirements_and_resume['requirements_id'],
                    "resume_id": create_requirements_and_resume['resume_id'],
                    "callback_url": callback_url
                },
                headers={"Authorization": f"Bearer {create_requirements_and_resume['access_token']}"}
            )
            assert response.status_code == 200
            # ответ от сервера проверять не надо, там всегда будут одни и те же данные, если получаем код 200

            # Данные с Kafka
            data_kafka = None
            for i in range(40):
                try:
                    msg = consumer.poll(timeout=1.0)
                    if msg is None:
                        if i == 39:
                            raise Exception("Не удалось получить сообщение от Kafka!")
                        continue
                    if msg.key().decode('utf-8') == KEY_NEW_REQUEST:
                        data_kafka = json.loads(msg.value().decode('utf-8'))
                    else:
                        raise Exception(f"Ожидался ключ '{KEY_NEW_REQUEST}', получен {msg.key().decode('utf-8')}")
                    break
                except KafkaError as e:
                    raise Exception(f"Ошибка Kafka: {e}")

            assert data_kafka['user_id'] == create_requirements_and_resume['user_id']
            assert data_kafka['requirements_id'] == create_requirements_and_resume['requirements_id']
            assert data_kafka['resume_id'] == create_requirements_and_resume['resume_id']
            assert data_kafka['callback_url'] == callback_url
            assert data_kafka['requirements'] == create_requirements_and_resume['requirements']
            assert data_kafka['resume'] == create_requirements_and_resume['resume']

    @pytest.mark.asyncio
    async def test_code_403(self, db_session, create_requirements_and_resume, create_user):  # при попытке получить данные другого пользователя
        # Создаем второго пользователя
        another_user = User(user_id=create_user['user_id'] + 1)
        db_session.add(another_user)
        await db_session.commit()
        await db_session.refresh(another_user)

        # Создаем требования и резюме для второго пользователя
        another_requirements = Requirements(
            user_id=another_user.user_id,
            requirements='Требования другого пользователя'
        )
        another_resume = Resume(
            user_id=another_user.user_id,
            resume='Резюме другого пользователя'
        )
        db_session.add(another_requirements)
        db_session.add(another_resume)
        await db_session.commit()
        await db_session.refresh(another_requirements)
        await db_session.refresh(another_resume)

        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            # 1. Пытаемся использовать требования другого пользователя
            response = await ac.post(
                "/start_processing",
                json={
                    'requirements_id': another_requirements.requirements_id,
                    'resume_id': create_requirements_and_resume['resume_id'],
                },
                headers={"Authorization": f"Bearer {create_requirements_and_resume['access_token']}"}
            )
            assert response.status_code == 403

            # 2. Пытаемся использовать резюме другого пользователя
            response = await ac.post(
                "/start_processing",
                json={
                    'requirements_id': create_requirements_and_resume['requirements_id'],
                    'resume_id': another_resume.resume_id,
                },
                headers={"Authorization": f"Bearer {create_requirements_and_resume['access_token']}"}
            )
            assert response.status_code == 403

    @pytest.mark.asyncio
    @pytest.mark.parametrize(
        'data_request',
        [
            ({"requirements_id": 1, "resume_id": 423432}),
            ({"requirements_id": 543534, "resume_id": 1})
        ]
    )
    async def test_code_404(self, data_request, create_requirements_and_resume): # при вводе несуществующего ID
        async with AsyncClient(
                transport=ASGITransport(app),
                base_url="http://test",
        ) as ac:
            response = await ac.post(
                "/start_processing",
                json=data_request,
                headers={"Authorization": f"Bearer {create_requirements_and_resume['access_token']}"}
            )
            assert response.status_code == 404



